import os

from docx import Document
from docx.shared import Cm, Pt  # 引入cm单位，便于设置图片的宽度
from docx.enum.table import WD_TABLE_ALIGNMENT  # 用于设置单元格的内容居中对齐
from docx.table import _Cell
from docx.oxml.shared import OxmlElement, qn


def set_cell_margins(cell, **kwargs):
    '''设置某单元格间距
    长度单位为Twips，1Twips = 1/20pt，1Twips = 1/567cm
    :param cell: 某单元格
    :param top: 上边距
    :param start: 左边距
    :param bottom: 下边距
    :param end: 右边距
    '''
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m in ['top', 'start', 'bottom', 'end']:
        if m in kwargs:
            node = OxmlElement('w:{}'.format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)


def add_pic_to_cell(table, row, col, pic_path, pic_wid):
    cell = table.cell(row, col)  # 获取某单元格对象（从0开始索引）
    # cell.width = Pt(pic_wid).cm
    cell.width = Cm(pic_wid)
    
    set_cell_margins(cell, start=0, top=0, end=0, bottom=0)
    # 在单元格中添加段落，区块
    c_p1 = cell.paragraphs[0]
    c_p1.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置单元格内容居中对齐
    c_p1.paragraph_format.left_indent = Cm(0)
    c_p1.paragraph_format.right_indent = Cm(0)
    c_p1.paragraph_format.space_before = Cm(0)
    c_p1.paragraph_format.space_after = Cm(0)
    c_p1.paragraph_format.line_spacing = 1
    c_run1 = c_p1.add_run()
    # 在单元格中添加图片，我的图片是bar.png，图片和py文件在同一个目录下
    c_run1.add_picture(pic_path, width=Cm(pic_wid))


def add_text_to_cell(table, row, col, text, cell_wid):
    cell = table.cell(row, col)  # 获取某单元格对象（从0开始索引）
    # cell.width = Pt(cell_wid).cm
    cell.width = Cm(cell_wid)
    set_cell_margins(cell, start=0, top=0, end=0, bottom=0)
    # 在单元格中添加段落，区块
    c_p1 = cell.paragraphs[0]
    c_p1.paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER  # 设置单元格内容居中对齐
    c_p1.paragraph_format.left_indent = Cm(0)
    c_p1.paragraph_format.right_indent = Cm(0)
    c_p1.paragraph_format.space_before = Cm(0)
    c_p1.paragraph_format.space_after = Cm(0)
    c_run1 = c_p1.add_run()

    # 在单元格中添加图片，我的图片是bar.png，图片和py文件在同一个目录下
    c_run1.add_text(text)


def generate_table(imgs_dir, g_row, g_col, pic_wid, file_name, add_label=False, pic_labels=None):
    imgs = os.listdir(imgs_dir)
    # 创建文档
    document = Document()
    # document.add_heading('CVLab 论文图表生成 by xjx\n\n', level=1)
    # 添加表格
    if add_label:
        g_row = g_row * 2
    table = document.add_table(rows=g_row, cols=g_col)  # 添加一个1行1列的空表
    row_index = 0
    col_index = 0

    def sort_file(x):
        try:
            fname, ftype = os.path.splitext(x)

            return float(fname)
        except:
            return x

    imgs = sorted(imgs, key=sort_file)
    for idx, img in enumerate(imgs):
        img_path = os.path.join(imgs_dir, img)
        fname, _ = os.path.splitext(img)
        add_pic_to_cell(table, row_index, col_index, img_path, pic_wid)
        if add_label:
            if pic_labels is None:
                add_text_to_cell(table, row_index + 1, col_index, fname, pic_wid)
            else:
                add_text_to_cell(table, row_index + 1, col_index, pic_labels[idx], pic_wid)
            col_index = col_index + 1
            if col_index >= g_col:
                col_index = 0
                row_index = row_index + 2
        else:
            col_index = col_index + 1
            if col_index >= g_col:
                col_index = 0
                row_index = row_index + 1

    document.save(file_name)


# generate_table(图片路径,行,列,图片大小(CM),输出文件,是否生成图像下标签,自定义标签)
if __name__ == '__main__':
    generate_table("imgs", 10, 4, 3.75, "out2.docx", False)
